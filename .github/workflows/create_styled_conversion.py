#!/usr/bin/env python3
"""
Convert markdown with Mermaid diagrams to .docx with custom styling.
Uses python-docx to create a properly styled reference document.

Usage:
    python3 create_styled_conversion.py <input_markdown_file> [output_docx_file]

Arguments:
    input_markdown_file: Path to the input markdown file (required)
    output_docx_file: Path to the output DOCX file (optional, defaults to same name as input with .docx extension)
"""

import re
import subprocess
import os
import sys
import argparse
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_reference_docx():
    """Create a reference .docx with custom styling for code blocks and tables."""
    doc = Document()

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Modify Source Code style for code blocks
    styles = doc.styles

    # Get or create Source Code style
    try:
        code_style = styles['Source Code']
    except KeyError:
        code_style = styles.add_style('Source Code', 1)  # 1 = paragraph style

    # Style the code blocks
    code_font = code_style.font
    code_font.name = 'Consolas'
    code_font.size = Pt(9)
    code_font.color.rgb = RGBColor(51, 51, 51)

    # Set paragraph formatting
    code_para = code_style.paragraph_format
    code_para.left_indent = Inches(0.25)
    code_para.right_indent = Inches(0.25)
    code_para.space_before = Pt(6)
    code_para.space_after = Pt(6)

    # Background shading for code blocks
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F0F0')  # Light gray background
    code_style._element.get_or_add_pPr().append(shd)

    # Get or create Verbatim Char style for inline code
    try:
        inline_code_style = styles['Verbatim Char']
    except KeyError:
        inline_code_style = styles.add_style('Verbatim Char', 2)  # 2 = character style

    inline_font = inline_code_style.font
    inline_font.name = 'Consolas'
    inline_font.size = Pt(9)
    inline_font.color.rgb = RGBColor(199, 37, 78)  # Reddish color for inline code

    # Create a sample table with proper styling
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'  # Use Table Grid style which has borders

    # Add sample content to table
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Header 1'
    header_cells[1].text = 'Header 2'
    header_cells[2].text = 'Header 3'

    # Style header row
    for cell in header_cells:
        # Make header bold
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        # Add gray background to header
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')  # Light gray for headers
        cell._element.get_or_add_tcPr().append(shading_elm)

    # Add data rows
    row1_cells = table.rows[1].cells
    row1_cells[0].text = 'Data 1'
    row1_cells[1].text = 'Data 2'
    row1_cells[2].text = 'Data 3'

    row2_cells = table.rows[2].cells
    row2_cells[0].text = 'Data 4'
    row2_cells[1].text = 'Data 5'
    row2_cells[2].text = 'Data 6'

    # Save reference document
    doc.save('custom-reference.docx')
    print("✓ Created custom reference document: custom-reference.docx")
    print("  - Code blocks: Gray background (F0F0F0)")
    print("  - Inline code: Reddish color")
    print("  - Tables: Grid borders with gray headers")

def extract_mermaid_blocks(markdown_content):
    """Extract all mermaid code blocks from markdown."""
    pattern = r'```mermaid\n(.*?)```'
    matches = re.finditer(pattern, markdown_content, re.DOTALL)
    return [(match.group(1), match.start(), match.end()) for match in matches]

def convert_mermaid_to_image(mermaid_code, output_path):
    """
    Check if mermaid diagram image already exists.
    Conversion is now handled by Docker container in the workflow.
    """
    return Path(output_path).exists()

def process_markdown_with_mermaid(input_file, image_path_prefix=''):
    """Process markdown file and replace mermaid blocks with image references.

    Args:
        input_file: Path to the input markdown file
        image_path_prefix: Optional prefix to prepend to image paths (e.g., '.github/workflows/')
    """

    # Read the markdown file
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Extract mermaid blocks
    mermaid_blocks = extract_mermaid_blocks(content)
    print(f"Found {len(mermaid_blocks)} mermaid diagrams")

    # Create a directory for images
    img_dir = Path('mermaid-diagrams')
    img_dir.mkdir(exist_ok=True)

    # Convert each mermaid block to an image and replace in markdown
    modified_content = content
    offset = 0

    for idx, (mermaid_code, start, end) in enumerate(mermaid_blocks):
        img_filename = f"diagram_{idx + 1}.png"
        img_path = img_dir / img_filename

        # Images should already be converted
        success = convert_mermaid_to_image(mermaid_code, str(img_path))
        if success:
            print(f"Diagram {idx + 1}: Using pre-rendered image")
        else:
            print(f"  ✗ Warning: Diagram {idx + 1} image not found at {img_path}")

        if success:
            # Replace mermaid block with image reference
            if image_path_prefix:
                # Use provided prefix (e.g., for GitHub Actions)
                full_path = f"{image_path_prefix.rstrip('/')}/{img_path}"
                image_markdown = f"\n![Diagram {idx + 1}]({full_path})\n"
            else:
                # Use relative path (for local builds)
                image_markdown = f"\n![Diagram {idx + 1}]({img_path})\n"

            # Adjust positions based on previous replacements
            adjusted_start = start + offset
            adjusted_end = end + offset

            modified_content = (
                modified_content[:adjusted_start] +
                image_markdown +
                modified_content[adjusted_end:]
            )

            # Update offset for next replacement
            offset += len(image_markdown) - (end - start)

    # Write modified markdown to temporary file
    temp_md = 'temp_with_images.md'
    with open(temp_md, 'w', encoding='utf-8') as f:
        f.write(modified_content)

    print(f"\n✓ Created temporary markdown file: {temp_md}")
    return temp_md

if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Convert markdown with Mermaid diagrams to DOCX with custom styling.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python3 create_styled_conversion.py input.md
  python3 create_styled_conversion.py input.md output.docx
  python3 create_styled_conversion.py ../../GITHUB_POLICIES_AND_GUIDELINES.md
        """
    )

    parser.add_argument(
        'input_file',
        help='Path to the input markdown file'
    )

    parser.add_argument(
        '--image-path-prefix',
        default='',
        help='Prefix to prepend to image paths (e.g., ".github/workflows/" for GitHub Actions)'
    )

    args = parser.parse_args()

    # Validate input file exists
    input_path = Path(args.input_file)
    if not input_path.exists():
        print(f"Error: Input file not found: {args.input_file}")
        sys.exit(1)

    print(f"Input file: {args.input_file}")
    if args.image_path_prefix:
        print(f"Image path prefix: {args.image_path_prefix}")

    # Create reference document with custom styling
    print("\nCreating custom reference document...")
    create_reference_docx()

    print("\nProcessing markdown and replacing mermaid diagrams with images...")
    temp_md = process_markdown_with_mermaid(args.input_file, args.image_path_prefix)

    print("\n✓ Markdown processing complete!")
    print(f"  - Temporary markdown: {temp_md}")
    print(f"  - Reference document: custom-reference.docx")
    print(f"  - Next: Run pandoc to convert to DOCX")
