#!/usr/bin/env python3
"""
Add borders to all tables in a .docx file for better readability.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#000000"},
        bottom={"sz": 12, "val": "single", "color": "#000000"},
        start={"sz": 12, "val": "single", "color": "#000000"},
        end={"sz": 12, "val": "single", "color": "#000000"},
    )
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Remove existing borders
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tcPr.remove(tcBorders)

    # Create new borders element
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'start', 'bottom', 'end'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            for key, value in edge_data.items():
                edge_el.set(qn(f'w:{key}'), str(value))
            tcBorders.append(edge_el)

    tcPr.append(tcBorders)

def add_table_borders(docx_path, output_path=None):
    """Add borders to all tables in the document."""
    if output_path is None:
        output_path = docx_path

    doc = Document(docx_path)

    # Border settings
    border_settings = {
        "sz": "12",  # Border size (1/8 pt, so 12 = 1.5pt)
        "val": "single",  # Border style
        "color": "000000"  # Black color
    }

    header_shading = "D9D9D9"  # Light gray for headers

    print(f"Processing {len(doc.tables)} tables...")

    for table_idx, table in enumerate(doc.tables):
        print(f"  Table {table_idx + 1}: {len(table.rows)} rows × {len(table.columns)} columns")

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                # Add borders to all cells
                set_cell_border(
                    cell,
                    top=border_settings,
                    bottom=border_settings,
                    start=border_settings,
                    end=border_settings
                )

                # Style header row (first row)
                if row_idx == 0:
                    # Add gray background
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), header_shading)
                    cell._element.get_or_add_tcPr().append(shading_elm)

                    # Make text bold
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

    doc.save(output_path)
    print(f"✓ Saved document with table borders: {output_path}")

if __name__ == '__main__':
    input_file = 'GITHUB_POLICIES_AND_GUIDELINES.docx'

    print("Adding borders to all tables...")
    add_table_borders(input_file)
    print("\n✓ All tables now have borders!")
    print("  - Black borders on all cells")
    print("  - Gray background on header rows")
    print("  - Bold text in headers")